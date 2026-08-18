import logging
import os
import uuid

try:
    from ..domain.entities import LegalDocument
except Exception:
    try:
        from domain.entities import LegalDocument
    except Exception:
        LegalDocument = None
try:
    from ..infrastructure import storage, db, parser
except Exception:
    try:
        from infrastructure import storage, db, parser
    except Exception:
        storage = db = parser = None
from .interfaces.doc_converter import DocConverter
try:
    from ..infrastructure.doc_converter.libreoffice_converter import LibreOfficeConverter
except Exception:
    try:
        from infrastructure.doc_converter.libreoffice_converter import LibreOfficeConverter
    except Exception:
        LibreOfficeConverter = None

logger = logging.getLogger(__name__)


def handle_upload(fileobj, filename, title=None, reference_number=None, issued_date=None, uploader_id=None, content_hash=None):
    doc_id = str(uuid.uuid4())
    # Read raw bytes to compute content hash and ensure storage receives a fresh BytesIO
    try:
        # fileobj may be an UploadFile/file-like or BytesIO
        fileobj.seek(0)
    except Exception:
        pass
    data = None
    try:
        data = fileobj.read()
    except Exception:
        # fallback: try reading from .file attribute
        try:
            data = getattr(fileobj, 'file').read()
        except Exception:
            data = None

    import hashlib, io
    if data is None:
        # if we couldn't read, delegate to storage directly (best-effort)
        storage_path = storage.save_file(fileobj, doc_id, filename)
        computed_hash = None
    else:
        computed_hash = hashlib.sha256(data).hexdigest()
        storage_path = storage.save_file(io.BytesIO(data), doc_id, filename)
    record = {
        'id': doc_id,
        'title': title,
        'reference_number': reference_number,
        'issued_date': issued_date,
        'original_filename': filename,
        'storage_path': storage_path,
        'uploader_id': uploader_id,
        'status': 'PENDING',
        'content_hash': content_hash or computed_hash,
        'extraction_status': 'EXTRACTION_PENDING',
    }
    db.insert_document(record)
    return doc_id


def _move_to_failed_storage(path, error_message):
    moved = storage.move_to_failed(path)
    if moved:
        logger.error('Moved failed document artifact to failed storage: %s', moved)
    return moved


def run_parsing_by_id(document_id):
    document_record = db.get_document(document_id)
    if not document_record:
        logger.warning('No document found for parsing: %s', document_id)
        return
    try:
        logger.info('run_parsing_by_id: starting parse for document %s', document_id)
        db.update_status(document_id, 'PROCESSING')
        run_parsing_for_document(dict(document_record))
        logger.info('run_parsing_by_id: finished parse for document %s', document_id)
    except Exception as exc:
        logger.exception('Parsing failed for document %s', document_id)
        db.update_status(document_id, 'FAILED', error_message=str(exc))
        raise


def run_parsing_for_document(document_record):
    if not document_record:
        raise ValueError('document_record is required')

    document_id = document_record['id']
    path = document_record['storage_path']
    temp_paths = []
    extracted_path = None
    try:
        # PDF-first workflow: require PDF input and extract first-page lines
        if not path.lower().endswith('.pdf'):
            # try converting .doc/.docx to docx using LibreOffice then extract paragraphs as lines
            if path.lower().endswith('.doc') or path.lower().endswith('.docx'):
                if LibreOfficeConverter is None:
                    raise RuntimeError('Only PDF input supported and LibreOffice converter not available for doc/docx')
                conv = LibreOfficeConverter()
                # convert to docx if needed
                if path.lower().endswith('.doc'):
                    docx_path = conv.convert_to_docx(path)
                    temp_paths.append(docx_path)
                    path = docx_path
                # if we have a docx file, extract paragraphs and map to lines
                try:
                    from docx import Document
                    doc = Document(path)
                    lines = [l for l in (p.text for p in doc.paragraphs) if l and l.strip()]
                except Exception as exc:
                    logger.exception('Failed to extract lines from docx fallback for %s: %s', path, exc)
                    db.update_extraction_metadata(document_id, extraction_status='EXTRACTION_FAILED', extraction_error=str(exc))
                    raise
            else:
                raise RuntimeError('Only PDF/docx documents are supported by the PDF-centric pipeline')
        else:
            # extract lines from FULL PDF (all pages) for clause parsing
            try:
                # prefer full-document extraction for structural parsing
                lines = parser.extract_all_lines_from_pdf(path)
            except Exception as exc:
                logger.exception('Failed to extract lines from PDF %s: %s', path, exc)
                db.update_extraction_metadata(document_id, extraction_status='EXTRACTION_FAILED', extraction_error=str(exc))
                raise

        # Save extracted artifact as list of lines for traceability
        items = [{'type': 'line', 'text': l} for l in lines]
        logger.info('run_parsing_for_document: document=%s path=%s total_lines=%d', document_id, path, len(lines))
        extracted_path = None
        try:
            extracted_path = storage.save_extracted_artifact(document_id, items)
            db.update_extraction_metadata(document_id, extracted_path=extracted_path, extraction_status='EXTRACTION_PENDING')
        except Exception as exc:
            logger.error('Failed to save extracted artifact for document %s: %s', document_id, exc)
            db.update_extraction_metadata(document_id, extraction_status='EXTRACTION_FAILED', extraction_error=f'failed to save extracted artifact: {exc}')
            raise

        metadata = parser.extract_document_metadata(
            lines,
            fallback_title=os.path.splitext(document_record.get('original_filename', 'document'))[0],
        )

        # If parser did not find a reference number, fallback to filename-based pattern
        if not metadata.get('reference_number'):
            orig = document_record.get('original_filename', '')
            base = os.path.splitext(orig)[0]
            if base and '.' in base:
                fallback_ref = '/'.join(base.split('.'))
                metadata['reference_number'] = fallback_ref

        logger.debug('Extracted metadata - title=%s, ref=%s, date=%s', 
                    metadata.get('title'), metadata.get('reference_number'), metadata.get('issued_date'))

        db.update_document_metadata(
            document_id,
            title=metadata.get('title'),
            reference_number=metadata.get('reference_number'),
            issued_date=metadata.get('issued_date'),
        )

        # update extraction status depending on success
        if metadata.get('title') and metadata.get('reference_number'):
            db.update_extraction_metadata(document_id, extracted_path=extracted_path, extraction_status='SUCCESS', extraction_error='')
            logger.info('Extraction successful for document %s', document_id)
        else:
            db.update_extraction_metadata(document_id, extracted_path=extracted_path, extraction_status='NEEDS_REVIEW', extraction_error='missing title or reference_number')
            logger.warning('Extraction incomplete for document %s, marked NEEDS_REVIEW', document_id)

        ch = document_record.get('content_hash')
        if not ch:
            import hashlib
            with open(path, 'rb') as fh:
                data = fh.read()
            content_hash = hashlib.sha256(data).hexdigest()
            db.update_extraction_metadata(document_id, content_hash=content_hash)

        # For compatibility, still run paragraph parsing on full text if available
        nodes = []
        try:
            nodes = parser.parse_paragraphs(lines)
        except Exception:
            # don't crash the whole run for parsing node issues, but bubble up the exception
            logger.exception('Failed to parse paragraphs for document %s', document_id)
            raise
        inserted = 0
        if nodes:
            db.insert_clauses(document_id, nodes)
            inserted = len(nodes)
        logger.info('run_parsing_for_document: document=%s inserted_clauses=%d', document_id, inserted)
        
        # Note: extraction_status already updated above based on metadata

        db.update_status(document_id, 'PROCESSED')
        logger.info('Parsing completed for document %s: %d clause(s)', document_id, len(nodes or []))
    except Exception as exc:
        logger.exception('Processing failed for document %s', document_id)
        _move_to_failed_storage(path, str(exc))
        db.update_status(document_id, 'FAILED', error_message=str(exc))
        raise
    finally:
        # Cleanup temporary scratch files after each processing step, while preserving canonical source/extracted output.
        storage.cleanup_temp_files(temp_paths)


def enhance_metadata_with_gemini(document_id: str):
    """Background task: call AI to enhance metadata and persist results.

    Uses parser.extract_metadata_via_ai() which wraps the Gemini/genai client.
    Updates document metadata and extraction_status accordingly.
    """
    document_record = db.get_document(document_id)
    if not document_record:
        logger.warning('enhance_metadata_with_gemini: document not found: %s', document_id)
        return
    path = document_record.get('storage_path')
    fallback_title = os.path.splitext(document_record.get('original_filename', 'document'))[0]
    try:
        text_content = ''
        if path and path.lower().endswith('.pdf'):
            try:
                lines = parser.extract_lines_from_pdf(path)
                text_content = '\n'.join(lines[:50])
            except Exception as exc:
                logger.exception('AI enhancement: failed to extract PDF lines for %s: %s', path, exc)
                text_content = ''
        else:
            # try docx paragraphs
            try:
                from docx import Document
                doc = Document(path)
                lines = [p.text for p in doc.paragraphs if p and p.text]
                text_content = '\n'.join(lines[:50])
            except Exception:
                # try converting via LibreOffice if available
                if LibreOfficeConverter is not None:
                    try:
                        conv = LibreOfficeConverter()
                        if path.lower().endswith('.doc'):
                            docx_path = conv.convert_to_docx(path)
                        else:
                            docx_path = path
                        from docx import Document
                        doc = Document(docx_path)
                        lines = [p.text for p in doc.paragraphs if p and p.text]
                        text_content = '\n'.join(lines[:50])
                    except Exception as exc:
                        logger.exception('AI enhancement: failed to convert/extract doc for %s: %s', path, exc)
                        text_content = ''
                else:
                    text_content = ''

        if not text_content:
            logger.warning('enhance_metadata_with_gemini: no text content for %s', document_id)
            db.update_extraction_metadata(document_id, extraction_status='FAILED', extraction_error='no text for AI')
            return

        ai_meta = parser.extract_metadata_via_ai(text_content)
        if not ai_meta:
            logger.warning('enhance_metadata_with_gemini: AI returned no metadata for %s', document_id)
            db.update_extraction_metadata(document_id, extraction_status='FAILED', extraction_error='ai returned no metadata')
            return

        # persist updated metadata
        db.update_document_metadata(document_id, title=ai_meta.get('title'), reference_number=ai_meta.get('reference_number'), issued_date=ai_meta.get('issued_date'))
        db.update_extraction_metadata(document_id, extraction_status='SUCCESS')
        logger.info('enhance_metadata_with_gemini: updated metadata for %s', document_id)
    except Exception as exc:
        logger.exception('AI enhancement failed for %s: %s', document_id, exc)
        try:
            db.update_extraction_metadata(document_id, extraction_status='FAILED', extraction_error=str(exc))
        except Exception:
            pass
