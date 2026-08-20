from fastapi import FastAPI, UploadFile, File, Depends, HTTPException, BackgroundTasks, status, Query
from fastapi.openapi.utils import get_openapi
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from uuid import UUID
from ..application import upload_service
from ..infrastructure import db, storage
from ..infrastructure.auth.jwt_service import JWTService
from ..infrastructure import parser
from ..infrastructure.doc_converter.libreoffice_converter import LibreOfficeConverter
import os
import tempfile
from io import BytesIO
from docx import Document
from typing import Annotated, List
import re


ALLOWED_EXT = {'.doc', '.docx', '.pdf'}
ALLOWED_MIME_TYPES = {
    'application/msword',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'application/pdf',
    # Swagger UI / browsers may send generic or alternative MIME types
    'application/octet-stream',
    'application/x-pdf',
}
MAX_SIZE_BYTES = int(os.getenv('MAX_UPLOAD_BYTES', 20 * 1024 * 1024))

app = FastAPI(title='Legal Uploader Service')


def _fix_openapi_file_items(schema: dict) -> dict:
    # Ensure file upload item schemas include format: binary so Swagger UI
    # renders a file input instead of a plain string input.
    try:
        comps = schema.get('components', {}).get('schemas', {})
        for name, s in comps.items():
            props = s.get('properties', {})
            for pname, p in props.items():
                if isinstance(p, dict) and p.get('type') == 'array' and 'items' in p:
                    item = p['items']
                    if isinstance(item, dict) and item.get('type') == 'string':
                        # set format to binary if not present
                        if 'format' not in item:
                            item['format'] = 'binary'
                        # prefer application/pdf when a contentMediaType is present
                        if 'contentMediaType' in item and 'pdf' in item.get('contentMediaType', ''):
                            item['contentMediaType'] = 'application/pdf'
                        # write back
                        p['items'] = item
    except Exception:
        pass
    return schema


def _custom_openapi():
    if app.openapi_schema:
        return app.openapi_schema
    openapi_schema = get_openapi(title=app.title, version="0.1.0", routes=app.routes)
    openapi_schema = _fix_openapi_file_items(openapi_schema)
    app.openapi_schema = openapi_schema
    return app.openapi_schema


app.openapi = _custom_openapi

security = HTTPBearer()
jwt_service = JWTService()


def allowed_file(filename: str) -> bool:
    ext = os.path.splitext(filename)[1].lower()
    return ext in ALLOWED_EXT


def is_allowed_mime(ctype: str) -> bool:
    """Return True for known allowed MIME types. Accepts MIME with params
    (e.g. 'application/pdf; charset=binary') and common pdf variants.
    """
    if not ctype:
        return False
    base = ctype.split(';', 1)[0].strip().lower()
    # accept any MIME that explicitly mentions pdf or is in the allowlist
    if base in ALLOWED_MIME_TYPES:
        return True
    if 'pdf' in base:
        return True
    return False


def normalize_uuid(value):
    if value is None:
        return None
    if isinstance(value, str):
        try:
            return str(UUID(value))
        except ValueError:
            return None
    return str(value)


def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    token = credentials.credentials
    payload = jwt_service.verify_token(token)
    if not payload:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail='Invalid token')
    return payload


def extract_from_docx_bytes(file_bytes: bytes, fallback_title: str, filename: str = 'document.docx') -> dict:
    ext = os.path.splitext(filename)[1].lower()
    if ext == '.doc':
        with tempfile.TemporaryDirectory() as tmpdir:
            temp_path = os.path.join(tmpdir, filename)
            with open(temp_path, 'wb') as fh:
                fh.write(file_bytes)
            converted = LibreOfficeConverter().convert_to_docx(temp_path)
            with open(converted, 'rb') as fh:
                file_bytes = fh.read()
    doc = Document(BytesIO(file_bytes))
    lines = [p.text for p in doc.paragraphs[:10]]
    return parser.extract_document_metadata(lines, fallback_title=fallback_title)


@app.post('/api/laws/upload')
async def upload_law(
    files: Annotated[List[UploadFile], File(description="Chọn 1 hoặc nhiều file PDF")],
    user=Depends(get_current_user),
    background_tasks: BackgroundTasks = None,
    force: bool = Query(False),
):
    if not files or len(files) == 0:
        raise HTTPException(status_code=400, detail='files list is required')

    # Read and validate files
    entries = []  # list of (filename, bytes, content_type)
    for f in files:
        fname = f.filename or ''
        ctype = getattr(f, 'content_type', None)
        if not allowed_file(fname):
            if ctype and is_allowed_mime(ctype):
                pass
            else:
                raise HTTPException(status_code=415, detail=f'unsupported media type for {fname}')
        contents = await f.read()
        if len(contents) > MAX_SIZE_BYTES:
            raise HTTPException(status_code=413, detail=f'payload too large for file {fname}')
        entries.append((fname, contents, ctype))

    # If single file, keep behavior identical
    if len(entries) == 1:
        filename, contents, ctype = entries[0]
    else:
        # sort by numeric token in filename if available (e.g., trang-1.pdf, page_2.pdf)
        def extract_num(s: str):
            if not s:
                return None
            m = re.search(r"(\d+)(?=[^\d]*$)", s)
            if m:
                try:
                    return int(m.group(1))
                except Exception:
                    return None
            return None

        entries = sorted(entries, key=lambda e: (extract_num(e[0]) if extract_num(e[0]) is not None else float('inf'), e[0].lower()))

        # Merge PDFs in-memory using pypdf (PdfMerger or PdfReader/PdfWriter)
        try:
            from pypdf import PdfMerger  # type: ignore
            merger_available = True
        except Exception:
            merger_available = False

        try:
            if merger_available:
                from pypdf import PdfMerger
                merger = PdfMerger()
                try:
                    for fn, b, ct in entries:
                        merger.append(BytesIO(b))
                    merged_io = BytesIO()
                    merger.write(merged_io)
                    merged_io.seek(0)
                    contents = merged_io.read()
                finally:
                    try:
                        merger.close()
                    except Exception:
                        pass
            else:
                from pypdf import PdfReader, PdfWriter
                writer = PdfWriter()
                for fn, b, ct in entries:
                    reader = PdfReader(BytesIO(b))
                    for p in reader.pages:
                        writer.add_page(p)
                merged_io = BytesIO()
                writer.write(merged_io)
                merged_io.seek(0)
                contents = merged_io.read()
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f'failed to merge PDF parts: {exc}')

        filename = 'merged.pdf'
        ctype = 'application/pdf'

    # MIME/basic checks for final contents
    if not contents:
        raise HTTPException(status_code=400, detail='no file contents provided')

    # compute SHA256 content hash
    import hashlib
    content_hash = hashlib.sha256(contents).hexdigest()

    fallback_title = os.path.splitext(filename)[0]
    # Extract metadata depending on file type: PDF vs DOC/DOCX
    try:
        ext = os.path.splitext(filename)[1].lower()
        if ext == '.pdf' or (not ext and (ctype and 'pdf' in (ctype or '').lower())):
            # write to a temp file and let parser handle PDF extraction
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=True) as tf:
                tf.write(contents)
                tf.flush()
                try:
                    lines = parser.extract_lines_from_pdf(tf.name)
                except Exception as e:
                    raise RuntimeError(f'failed to extract PDF lines: {e}') from e
            metadata = parser.extract_document_metadata(lines, fallback_title=fallback_title)
        else:
            # doc/docx flow
            metadata = extract_from_docx_bytes(contents, fallback_title=fallback_title, filename=filename)
    except (ValueError, RuntimeError) as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f'internal error during metadata extraction: {e}')
    title = metadata.get('title') or fallback_title
    reference_number = metadata.get('reference_number')
    issued_date = metadata.get('issued_date')
    uploader_id = normalize_uuid(user.get('sub')) if isinstance(user, dict) else None

    fileobj = BytesIO(contents)
    try:
        # Detect duplicate
        existing = db.find_document_by_hash(content_hash)
        if existing and not force:
            return {'id': str(existing['id']), 'status': existing.get('status'), 'message': 'DUPLICATE', 'content_hash': content_hash}

        # Persist the uploaded file and initial record
        doc_id = upload_service.handle_upload(fileobj, filename, title, reference_number, issued_date, uploader_id, content_hash=content_hash)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    # Always schedule full parsing of the stored document (clause extraction) asynchronously
    try:
        if background_tasks is not None:
            background_tasks.add_task(upload_service.run_parsing_by_id, doc_id)
        else:
            # synchronous fallback (best-effort)
            try:
                upload_service.run_parsing_by_id(doc_id)
            except Exception:
                pass
    except Exception:
        # scheduling failure should not block upload response
        pass

    # Fast-path validation: if the rule-based metadata is sufficient, mark SUCCESS;
    # otherwise schedule background AI enhancement and return immediately.
    def validate_metadata(m: dict) -> bool:
        import re
        is_valid = True
        if not m or not isinstance(m, dict):
            return False
        title = (m.get('title') or '').strip()
        reference_number = (m.get('reference_number') or '').strip()

        # title must exist and be reasonably long
        if not title or len(title) < 5:
            return False

        # reject titles that look like header/boilerplate containing these keywords
        bad_keywords = [
            'cộng hòa',
            'xã hội chủ nghĩa',
            'độc lập',
            'tự do',
            'công báo',
            'quốc hội',
        ]
        low_title = title.lower()
        for kw in bad_keywords:
            if kw in low_title:
                return False

        # reference number must exist
        if not reference_number:
            return False

        # if reference contains a date-like token (DD-MM-YYYY or DD/MM/YYYY) -> invalid
        if re.search(r"\b\d{1,2}[-/]\d{1,2}[-/]\d{4}\b", reference_number):
            return False

        # if reference_number does not contain any slash '/' -> invalid
        if '/' not in reference_number:
            return False

        return is_valid

    # ensure we have `lines` for saving extracted artifact
    lines_for_artifact = []
    try:
        ext = os.path.splitext(filename)[1].lower()
        if ext == '.pdf' or (not ext and (ctype and 'pdf' in (ctype or '').lower())):
            # already extracted `lines` above when computing metadata for PDFs
            # but ensure variable exists (it was defined earlier for PDF path)
            # if not, re-extract from stored file
            if 'lines' not in locals() or not lines:
                # get storage path from db
                rec = db.get_document(doc_id)
                if rec:
                    stored = rec.get('storage_path')
                    try:
                        lines = parser.extract_lines_from_pdf(stored)
                    except Exception:
                        lines = []
            lines_for_artifact = [{'type': 'line', 'text': l} for l in (lines or [])]
        else:
            # extract paragraphs from DOCX to produce artifact lines
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(BytesIO(contents))
                doc_lines = [p.text for p in doc.paragraphs if p and p.text]
                lines_for_artifact = [{'type': 'line', 'text': l} for l in doc_lines]
            except Exception:
                lines_for_artifact = []
    except Exception:
        lines_for_artifact = []

    # Update metadata fields in DB (best-effort)
    try:
        db.update_document_metadata(doc_id, title=title, reference_number=reference_number, issued_date=issued_date)
    except Exception:
        pass

    if validate_metadata(metadata):
        # Fast path success: save extracted artifact and mark done
        try:
            extracted_path = None
            if lines_for_artifact:
                extracted_path = storage.save_extracted_artifact(doc_id, lines_for_artifact)
            db.update_extraction_metadata(doc_id, extracted_path=extracted_path, extraction_status='SUCCESS')
            db.update_status(doc_id, 'PROCESSED')
        except Exception as e:
            # if update fails, surface as server error
            raise HTTPException(status_code=500, detail=f'failed to persist extraction metadata: {e}')
        return {'id': doc_id, 'status': 'PROCESSED', 'message': 'Uploaded and metadata extracted.'}
    else:
        # Slow path: persist as pending AI enhancement and schedule background task
        try:
            db.update_extraction_metadata(doc_id, extracted_path=None, extraction_status='PENDING_AI_ENHANCEMENT')
            db.update_status(doc_id, 'PROCESSED')
        except Exception:
            pass

        if background_tasks is not None:
            background_tasks.add_task(upload_service.enhance_metadata_with_gemini, doc_id)
        else:
            # best-effort synchronous fallback (not ideal in production)
            try:
                upload_service.enhance_metadata_with_gemini(doc_id)
            except Exception:
                pass

        return {'id': doc_id, 'status': 'PROCESSED', 'message': 'Uploaded. Metadata enhancement scheduled.'}


@app.patch('/api/laws/{document_id}/approve')
def approve_law(document_id: str, user=Depends(get_current_user)):
    doc = db.get_document(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail='not found')
    status_val = doc['status']
    if status_val not in ('DONE', 'PROCESSED'):
        raise HTTPException(status_code=400, detail=f'document must be in PROCESSED state to approve (current={status_val})')
    approver = normalize_uuid(user.get('sub')) if isinstance(user, dict) else None
    try:
        db.update_status(document_id, 'APPROVED', error_message=None, approved_by=approver)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return {'id': document_id, 'status': 'APPROVED'}

