#!/usr/bin/env python3
"""
Integration test: Simulate full upload flow and check DB writes.
"""
import sys
import os
import io
import uuid
from datetime import datetime

sys.path.insert(0, os.path.dirname(__file__))

# Override database host to use localhost instead of Docker hostname
os.environ['POSTGRES_HOST'] = 'localhost'

# Override DATABASE_URL to use localhost instead of Docker hostname
os.environ['DATABASE_URL'] = 'postgresql://demoxv_admin:demoxv_secret_pass@localhost:5432/legal_db'

# Create a minimal test docx file in memory
from docx import Document as DocxDocument

def create_test_docx():
    """Create a test .docx file with Vietnamese law metadata."""
    doc = DocxDocument()
    
    # Add paragraphs that simulate a Vietnamese law document header
    doc.add_paragraph("BỘ LUẬT DÂN SỰ")
    doc.add_paragraph("Số: 91/2015/QH13")
    doc.add_paragraph("ngày 24 tháng 11 năm 2015")
    doc.add_paragraph("")
    doc.add_paragraph("QUỐC HỘI BAN HÀNH")
    doc.add_paragraph("")
    doc.add_paragraph("Căn cứ Hiến pháp nước Cộng hòa Xã hội chủ nghĩa Việt Nam năm 1992 đã được sửa đổi, bổ sung;")
    doc.add_paragraph("")
    doc.add_paragraph("CHƯƠNG I")
    doc.add_paragraph("NHỮNG QUY ĐỊNH CHUNG")
    doc.add_paragraph("")
    doc.add_paragraph("Điều 1. Phạm vi điều chỉnh")
    
    # Save to BytesIO
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def test_integration():
    """Test full upload flow."""
    print("\n" + "="*70)
    print("INTEGRATION TEST: Full Upload Flow")
    print("="*70)
    
    # Create test document
    docx_file = create_test_docx()
    filename = f"test_law_91_2015_QH13.docx"
    
    # Step 1: Upload
    print(f"\n[STEP 1] Uploading {filename}...")
    from src.application.upload_service import handle_upload
    doc_id = handle_upload(docx_file, filename, uploader_id=str(uuid.uuid4()))
    print(f"  Document ID: {doc_id}")
    
    # Step 2: Check initial DB state
    print(f"\n[STEP 2] Checking initial DB state...")
    from src.infrastructure import db
    doc = db.get_document(doc_id)
    if doc:
        print(f"  title: {doc.get('title')}")
        print(f"  reference_number: {doc.get('reference_number')}")
        print(f"  issued_date: {doc.get('issued_date')}")
        print(f"  content_hash: {doc.get('content_hash')}")
        print(f"  extraction_status: {doc.get('extraction_status')}")
    else:
        print("  ERROR: Document not found in DB!")
        return False
    
    # Step 3: Run parsing
    print(f"\n[STEP 3] Running parsing...")
    try:
        from src.application.upload_service import run_parsing_for_document
        run_parsing_for_document(dict(doc))
        print(f"  Parsing completed")
    except Exception as exc:
        print(f"  ERROR during parsing: {exc}")
        import traceback
        traceback.print_exc()
        return False
    
    # Step 4: Check final DB state
    print(f"\n[STEP 4] Checking final DB state...")
    doc = db.get_document(doc_id)
    if doc:
        print(f"  title: {doc.get('title')}")
        print(f"  reference_number: {doc.get('reference_number')}")
        print(f"  issued_date: {doc.get('issued_date')}")
        print(f"  content_hash: {doc.get('content_hash')}")
        print(f"  extraction_status: {doc.get('extraction_status')}")
        print(f"  extracted_path: {doc.get('extracted_path')}")
        print(f"  status: {doc.get('status')}")
    else:
        print("  ERROR: Document not found in DB!")
        return False
    
    # Check if fields are populated correctly
    success = all([
        doc.get('title') == 'BỘ LUẬT DÂN SỰ',
        doc.get('reference_number') == '91/2015/QH13',
        doc.get('issued_date') and '2015-11-24' in str(doc.get('issued_date')),
        doc.get('content_hash'),
        doc.get('extraction_status') == 'SUCCESS',
    ])
    
    print(f"\n{'✓ PASS' if success else '✗ FAIL'}: Integration test {'passed' if success else 'failed'}")
    return success


if __name__ == '__main__':
    try:
        success = test_integration()
        sys.exit(0 if success else 1)
    except Exception as exc:
        print(f"\nFATAL ERROR: {exc}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
